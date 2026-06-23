"""
PDF resume + Word cover letter generator — 3 selectable themes.
Usage: python3 generate_application.py <folder> [--style executive|minimal|asymmetric]
Default style: executive
"""

import sys, os, re
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    BaseDocTemplate, Frame, PageTemplate, Paragraph,
    Spacer, HRFlowable, Table, TableStyle, KeepTogether, FrameBreak
)
from reportlab.platypus.flowables import Flowable
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W, H = letter

# ── Palette ───────────────────────────────────────────────────────────────────
C_NAVY   = colors.HexColor("#0A1628")
C_NAVY2  = colors.HexColor("#0D1F35")
C_TEAL   = colors.HexColor("#1A7A8A")
C_TEAL2  = colors.HexColor("#12616E")
C_LTEAL  = colors.HexColor("#E0F2F5")
C_GOLD   = colors.HexColor("#C9A84C")
C_GOLD2  = colors.HexColor("#E8C96A")
C_GRAY   = colors.HexColor("#1A202C")
C_LGRAY  = colors.HexColor("#718096")
C_RULE   = colors.HexColor("#CBD5E0")
C_WHITE  = colors.white
C_OFFWHT = colors.HexColor("#F7FAFC")
C_LITE   = colors.HexColor("#EDF2F7")
C_SGRAY  = colors.HexColor("#F0F4F8")   # sidebar bg for minimal theme


# ── Custom Flowables ──────────────────────────────────────────────────────────

class GoldBarHeading(Flowable):
    """Section heading: 3pt gold left bar + teal all-caps text (executive)."""
    def __init__(self, text, width, bar_color=None, text_color=None):
        super().__init__()
        self.text = text.upper()
        self._w = width
        self.bar_color  = bar_color  or C_GOLD
        self.text_color = text_color or C_TEAL
        self._h = 14

    def wrap(self, aw, ah):
        return self._w, self._h

    def draw(self):
        c = self.canv
        c.saveState()
        c.setFillColor(self.bar_color)
        c.rect(0, 1, 3.5, 10, fill=1, stroke=0)
        c.setFont("Helvetica-Bold", 8)
        c.setFillColor(self.text_color)
        c.drawString(8, 3, self.text)
        c.restoreState()


class SkillBar(Flowable):
    """Visual skill bar for sidebar (executive / asymmetric themes)."""
    def __init__(self, label, width, fill=0.82, bar_color=None):
        super().__init__()
        self.label = label
        self._w = width
        self.fill = fill
        self.bar_color = bar_color or C_TEAL
        self._h = 18

    def wrap(self, aw, ah):
        return self._w, self._h

    def draw(self):
        c = self.canv
        c.saveState()
        # Label
        c.setFont("Helvetica", 7.5)
        c.setFillColor(colors.HexColor("#CBD5E0"))
        c.drawString(0, 8, self.label)
        # Background track
        track_y = 3
        track_h = 3.5
        track_w = self._w - 2
        c.setFillColor(colors.HexColor("#1E3A55"))
        c.roundRect(0, track_y, track_w, track_h, 1.5, fill=1, stroke=0)
        # Fill bar
        c.setFillColor(self.bar_color)
        c.roundRect(0, track_y, track_w * self.fill, track_h, 1.5, fill=1, stroke=0)
        c.restoreState()


class SkillChip(Flowable):
    """Inline skill tag with rounded rect background (minimal theme sidebar)."""
    def __init__(self, skills_list, width, bg=None, text_color=None):
        super().__init__()
        self.skills = skills_list
        self._w = width
        self.bg = bg or C_LTEAL
        self.text_color = text_color or C_TEAL2
        # calculate height based on wrapping
        self._h = max(16, 12 * ((len(skills_list) // 3) + 1))

    def wrap(self, aw, ah):
        return self._w, self._h

    def draw(self):
        c = self.canv
        c.saveState()
        x, y = 0, self._h - 12
        for skill in self.skills:
            c.setFont("Helvetica", 6.8)
            sw = c.stringWidth(skill, "Helvetica", 6.8) + 8
            if x + sw > self._w:
                x = 0
                y -= 12
            c.setFillColor(self.bg)
            c.roundRect(x, y - 1, sw, 9, 2, fill=1, stroke=0)
            c.setFillColor(self.text_color)
            c.drawString(x + 4, y + 1, skill)
            x += sw + 4
        c.restoreState()


class TimelineDot(Flowable):
    """Teal circle for timeline (executive theme)."""
    def __init__(self, r=4):
        super().__init__()
        self.r = r

    def wrap(self, aw, ah):
        return self.r * 2, self.r * 2

    def draw(self):
        c = self.canv
        c.setFillColor(C_TEAL)
        c.circle(self.r, self.r, self.r, fill=1, stroke=0)


# ── Style helpers ─────────────────────────────────────────────────────────────
def S(name, **kw):
    d = dict(fontName="Helvetica", fontSize=9, leading=13,
             textColor=C_GRAY, spaceAfter=0, spaceBefore=0)
    d.update(kw)
    return ParagraphStyle(name, **d)

def strip_md(t):
    t = re.sub(r'\*\*(.*?)\*\*', r'\1', t)
    return re.sub(r'\*(.*?)\*', r'\1', t)

def rich(t):
    t = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', t)
    return re.sub(r'\*(.*?)\*', r'<i>\1</i>', t)


# ── Markdown parser ───────────────────────────────────────────────────────────
def parse_resume(md: str) -> dict:
    lines = md.strip().split("\n")
    data = dict(name="", subtitle="", contact="", sections=[])
    current_sec = None
    current_job = None

    for raw in lines:
        ln = raw.strip()
        if ln.startswith("# "):
            data["name"] = ln[2:].strip(); continue
        if ln.startswith("**") and ln.endswith("**") and not current_sec:
            data["subtitle"] = ln.strip("*"); continue
        if not current_sec and ("@" in ln or "(" in ln) and "|" in ln:
            data["contact"] = ln; continue
        if ln.startswith("## "):
            current_sec = {"name": ln[3:].strip(), "items": [], "jobs": []}
            data["sections"].append(current_sec)
            current_job = None; continue
        if current_sec is None: continue
        if ln.startswith("### "):
            parts = ln[4:].split("|")
            current_job = {
                "title":   parts[0].strip() if len(parts) > 0 else "",
                "company": parts[1].strip() if len(parts) > 1 else "",
                "dates":   parts[2].strip() if len(parts) > 2 else "",
                "bullets": []
            }
            current_sec["jobs"].append(current_job); continue
        if current_job and ln.startswith("- "):
            current_job["bullets"].append(ln[2:].strip()); continue
        current_sec["items"].append(raw)

    return data


# ══════════════════════════════════════════════════════════════════════════════
#  THEME 1 — DARK EXECUTIVE
# ══════════════════════════════════════════════════════════════════════════════
EXEC_HEADER_H = 1.1 * inch
EXEC_SIDEBAR  = 2.15 * inch
EXEC_GOLD_BAR = 2.5
EXEC_MX       = 0.28 * inch
EXEC_MY       = 0.28 * inch
EXEC_MAIN_X   = EXEC_SIDEBAR + EXEC_GOLD_BAR + 5
EXEC_MAIN_W   = W - EXEC_MAIN_X - EXEC_MX * 0.6
EXEC_CONTENT_H = H - EXEC_HEADER_H - EXEC_MY * 2


def draw_executive(canvas, doc):
    canvas.saveState()
    # Header: full navy
    canvas.setFillColor(C_NAVY)
    canvas.rect(0, H - EXEC_HEADER_H, W, EXEC_HEADER_H, fill=1, stroke=0)

    # Decorative overlapping circles — top-right of header
    for (cx, cy, r, alpha) in [
        (W - 0.45*inch, H - 0.05*inch, 0.70*inch, 0.13),
        (W - 0.95*inch, H - 0.20*inch, 0.52*inch, 0.09),
        (W - 0.20*inch, H - 0.65*inch, 0.38*inch, 0.07),
        (W - 1.30*inch, H - 0.55*inch, 0.28*inch, 0.06),
    ]:
        canvas.setFillColor(colors.Color(0.79, 0.66, 0.30, alpha=alpha))
        canvas.setStrokeColor(colors.Color(0.79, 0.66, 0.30, alpha=alpha * 1.8))
        canvas.setLineWidth(1.2)
        canvas.circle(cx, cy, r, fill=1, stroke=1)

    # Gold bottom rule on header
    canvas.setFillColor(C_GOLD)
    canvas.rect(0, H - EXEC_HEADER_H, W, 2.5, fill=1, stroke=0)

    # Left accent strip (full page height)
    canvas.rect(0, 0, 3.5, H - EXEC_HEADER_H, fill=1, stroke=0)

    # Sidebar background
    canvas.setFillColor(C_NAVY2)
    canvas.rect(3.5, 0, EXEC_SIDEBAR - 3.5, H - EXEC_HEADER_H, fill=1, stroke=0)

    # Gold divider line between sidebar and main
    canvas.setFillColor(C_GOLD)
    canvas.rect(EXEC_SIDEBAR, 0, EXEC_GOLD_BAR, H - EXEC_HEADER_H, fill=1, stroke=0)

    canvas.restoreState()


def exec_header_story(data):
    s_name = S("eN", fontName="Helvetica-Bold", fontSize=26,
                textColor=C_WHITE, leading=28, spaceAfter=2)
    s_sub  = S("eS", fontName="Helvetica", fontSize=9,
                textColor=C_GOLD, leading=12, spaceAfter=0)
    s_cont = S("eC", fontName="Helvetica", fontSize=7.8,
                textColor=colors.HexColor("#8EAAC8"), leading=11, alignment=TA_RIGHT)

    contact_parts = [p.strip() for p in data["contact"].split("|") if p.strip()]
    cont_str = "   ·   ".join(contact_parts)

    name_cell  = [Paragraph(data["name"], s_name)]
    if data["subtitle"]:
        name_cell.append(Paragraph(data["subtitle"], s_sub))
    cont_cell  = [Spacer(1, 10), Paragraph(cont_str, s_cont)]

    t = Table([[name_cell, cont_cell]],
              colWidths=[W * 0.60, W * 0.38])
    t.setStyle(TableStyle([
        ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",  (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING",   (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 0),
    ]))
    return [Spacer(1, 0.14*inch), t]


def exec_sidebar_story(data, compact=False):
    s_sec  = S("eSec", fontName="Helvetica-Bold", fontSize=6.8,
                textColor=C_GOLD, leading=9, spaceAfter=4, spaceBefore=12,
                letterSpacing=1.0)
    s_body = S("eBd",  fontName="Helvetica", fontSize=7.5,
                textColor=colors.HexColor("#A8C0D6"), leading=10.5, spaceAfter=0)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.16 * inch)]

    for sec in data["sections"]:
        sname = sec["name"]
        if sname not in SIDEBAR_SECTIONS:
            continue

        story.append(Paragraph(sname.upper(), s_sec))

        if sname in ("Core Skills", "Technical Skills", "Skills"):
            # Collect all skill strings
            skills = []
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    skills.append(strip_md(ls[2:]).strip())
                elif ls.startswith("|") and "---" not in ls:
                    cells = [c.strip() for c in ls.split("|")[1:-1]]
                    if cells and strip_md(cells[0]).lower() not in ("category", ""):
                        cat  = strip_md(cells[0])
                        vals = strip_md(cells[1]) if len(cells) > 1 else ""
                        # Add category as a "header" then skills as bars
                        story.append(Paragraph(f"<b>{cat}</b>", s_body))
                        for v in vals.split(","):
                            v = v.strip()
                            if v:
                                story.append(SkillBar(v, EXEC_SIDEBAR - 20))
                        continue
            for sk in skills:
                if len(sk) > 42: sk = sk[:40] + "…"
                story.append(SkillBar(sk, EXEC_SIDEBAR - 20))

        elif sname == "Education":
            for ln in sec["items"]:
                ls = ln.strip()
                if not ls or ls.startswith("|"): continue
                story.append(Paragraph(strip_md(ls), s_body))

        elif sname == "Certifications":
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    story.append(Paragraph(f"• {strip_md(ls[2:])}", s_body))

    return story


def exec_section_heading(text):
    # Inline heading — gold bar + teal text, no import of GoldBarHeading width here
    s = S("eH", fontName="Helvetica-Bold", fontSize=7.8,
           textColor=C_TEAL, leading=9, spaceBefore=9, spaceAfter=3,
           leftIndent=8)
    flowables = [GoldBarHeading(text, EXEC_MAIN_W), Spacer(1, 1)]
    return KeepTogether(flowables)


def exec_main_story(data, compact=False, max_bullets=None):
    s_job   = S("eJT", fontName="Helvetica-Bold", fontSize=10,
                 textColor=C_NAVY, leading=13)
    s_co    = S("eCo", fontName="Helvetica-Oblique", fontSize=8.5,
                 textColor=C_TEAL, leading=11, spaceAfter=1)
    s_dates = S("eDt", fontName="Helvetica", fontSize=8,
                 textColor=C_GOLD, leading=12, alignment=TA_RIGHT)
    s_bull  = S("eBu", fontName="Helvetica", fontSize=8.8,
                 textColor=C_GRAY, leading=12.5 if not compact else 11.5,
                 leftIndent=11, firstLineIndent=-8,
                 spaceAfter=2 if not compact else 1)
    s_sum   = S("eSu", fontName="Helvetica", fontSize=9,
                 textColor=C_GRAY, leading=13.5, spaceAfter=3)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.1 * inch)]
    job_idx = 0

    for sec in data["sections"]:
        if sec["name"] in SIDEBAR_SECTIONS: continue

        story.append(exec_section_heading(sec["name"]))

        # Text items (summary, table rows)
        table_rows = []
        for ln in sec["items"]:
            ls = ln.strip()
            if not ls: continue
            if ls.startswith("|") and "---" not in ls:
                table_rows.append(ls); continue
            if table_rows:
                story.append(_build_table(table_rows, EXEC_MAIN_W)); table_rows = []
            story.append(Paragraph(rich(ls), s_sum))
        if table_rows:
            story.append(_build_table(table_rows, EXEC_MAIN_W))

        for job in sec["jobs"]:
            cap = (max_bullets or {}).get(job_idx)
            left  = Paragraph(f"<b>{job['title']}</b>", s_job)
            right = Paragraph(job["dates"], s_dates)
            t = Table([[left, right]],
                      colWidths=[EXEC_MAIN_W * 0.70, EXEC_MAIN_W * 0.30])
            t.setStyle(TableStyle([
                ("VALIGN",       (0,0),(-1,-1),"BOTTOM"),
                ("LEFTPADDING",  (0,0),(-1,-1),0),
                ("RIGHTPADDING", (0,0),(-1,-1),0),
                ("TOPPADDING",   (0,0),(-1,-1),0),
                ("BOTTOMPADDING",(0,0),(-1,-1),1),
            ]))
            block = [t]
            if job["company"]:
                block.append(Paragraph(job["company"], s_co))
            bullets = job["bullets"] if cap is None else job["bullets"][:cap]
            for b in bullets:
                block.append(Paragraph(
                    f"<font color='#1A7A8A'>▸</font>  {rich(b)}", s_bull))
            story.append(KeepTogether(block[:2]))
            for item in block[2:]: story.append(item)
            story.append(Spacer(1, 3 if not compact else 1))
            job_idx += 1

    return story


def make_exec_frames():
    hf = Frame(0, H - EXEC_HEADER_H, W, EXEC_HEADER_H * 0.88,
               leftPadding=EXEC_MX, rightPadding=EXEC_MX * 0.3,
               topPadding=0, bottomPadding=0, id="hdr", showBoundary=0)
    sf = Frame(EXEC_MX * 0.3, EXEC_MY, EXEC_SIDEBAR - EXEC_MX * 0.3, EXEC_CONTENT_H,
               leftPadding=10, rightPadding=6,
               topPadding=0, bottomPadding=0, id="side", showBoundary=0)
    mf = Frame(EXEC_MAIN_X + EXEC_MX * 0.2, EXEC_MY, EXEC_MAIN_W, EXEC_CONTENT_H,
               leftPadding=3, rightPadding=3,
               topPadding=0, bottomPadding=0, id="main", showBoundary=0)
    return hf, sf, mf


# ══════════════════════════════════════════════════════════════════════════════
#  THEME 2 — MODERN MINIMAL
# ══════════════════════════════════════════════════════════════════════════════
MIN_HEADER_H  = 1.25 * inch
MIN_SIDEBAR   = 2.1  * inch
MIN_MX        = 0.32 * inch
MIN_MY        = 0.28 * inch
MIN_MAIN_X    = MIN_SIDEBAR + 1
MIN_MAIN_W    = W - MIN_MAIN_X - MIN_MX
MIN_CONTENT_H = H - MIN_HEADER_H - MIN_MY * 2


def draw_minimal(canvas, doc):
    canvas.saveState()
    # White page (default)
    canvas.setFillColor(C_OFFWHT)
    canvas.rect(0, 0, W, H, fill=1, stroke=0)
    # Header: white bg with navy name block on left
    canvas.setFillColor(C_WHITE)
    canvas.rect(0, H - MIN_HEADER_H, W, MIN_HEADER_H, fill=1, stroke=0)
    # Thick navy left accent bar in header
    canvas.setFillColor(C_NAVY)
    canvas.rect(0, H - MIN_HEADER_H, 6, MIN_HEADER_H, fill=1, stroke=0)
    # Teal rule at bottom of header
    canvas.setFillColor(C_TEAL)
    canvas.rect(0, H - MIN_HEADER_H, W, 3, fill=1, stroke=0)
    # Sidebar bg (light gray)
    canvas.setFillColor(C_SGRAY)
    canvas.rect(0, 0, MIN_SIDEBAR, H - MIN_HEADER_H, fill=1, stroke=0)
    # Teal thin rule between sidebar and main
    canvas.setFillColor(C_TEAL)
    canvas.rect(MIN_SIDEBAR, 0, 1.5, H - MIN_HEADER_H, fill=1, stroke=0)
    canvas.restoreState()


def min_header_story(data):
    s_name = S("mN", fontName="Helvetica-Bold", fontSize=28,
                textColor=C_NAVY, leading=30, spaceAfter=2)
    s_sub  = S("mS", fontName="Helvetica", fontSize=10,
                textColor=C_TEAL, leading=13, spaceAfter=0)
    s_cont = S("mC", fontName="Helvetica", fontSize=8,
                textColor=C_LGRAY, leading=12, alignment=TA_RIGHT)

    contact_parts = [p.strip() for p in data["contact"].split("|") if p.strip()]
    cont_str = "  ·  ".join(contact_parts)

    name_cell = [Paragraph(data["name"], s_name)]
    if data["subtitle"]:
        name_cell.append(Paragraph(data["subtitle"], s_sub))
    cont_cell = [Spacer(1, 14), Paragraph(cont_str, s_cont)]

    t = Table([[name_cell, cont_cell]],
              colWidths=[W * 0.58, W * 0.40])
    t.setStyle(TableStyle([
        ("VALIGN",       (0,0),(-1,-1),"MIDDLE"),
        ("LEFTPADDING",  (0,0),(-1,-1),0),
        ("RIGHTPADDING", (0,0),(-1,-1),0),
        ("TOPPADDING",   (0,0),(-1,-1),0),
        ("BOTTOMPADDING",(0,0),(-1,-1),0),
    ]))
    return [Spacer(1, 0.18*inch), t]


def min_sidebar_story(data, compact=False):
    s_sec  = S("mSec", fontName="Helvetica-Bold", fontSize=7.5,
                textColor=C_NAVY, leading=9, spaceAfter=5, spaceBefore=12)
    s_body = S("mBd",  fontName="Helvetica", fontSize=7.8,
                textColor=C_GRAY, leading=11, spaceAfter=1)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.14 * inch)]

    for sec in data["sections"]:
        sname = sec["name"]
        if sname not in SIDEBAR_SECTIONS: continue

        story.append(Paragraph(sname.upper(), s_sec))
        story.append(HRFlowable(width="100%", thickness=1.5, color=C_TEAL,
                                spaceAfter=4, spaceBefore=0))

        if sname in ("Core Skills", "Technical Skills", "Skills"):
            skills = []
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    skills.append(strip_md(ls[2:]).strip())
                elif ls.startswith("|") and "---" not in ls:
                    cells = [c.strip() for c in ls.split("|")[1:-1]]
                    if cells and strip_md(cells[0]).lower() not in ("category", ""):
                        cat  = strip_md(cells[0])
                        vals = strip_md(cells[1]) if len(cells) > 1 else ""
                        story.append(Paragraph(f"<b>{cat}</b>", s_body))
                        chips = [v.strip() for v in vals.split(",") if v.strip()]
                        if chips:
                            story.append(SkillChip(chips, MIN_SIDEBAR - 18,
                                                   bg=C_LTEAL, text_color=C_TEAL2))
                        continue
            if skills:
                story.append(SkillChip(skills, MIN_SIDEBAR - 18,
                                       bg=C_LTEAL, text_color=C_TEAL2))

        elif sname == "Education":
            for ln in sec["items"]:
                ls = ln.strip()
                if not ls or ls.startswith("|"): continue
                story.append(Paragraph(strip_md(ls), s_body))

        elif sname == "Certifications":
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    story.append(Paragraph(f"✓  {strip_md(ls[2:])}", s_body))

    return story


def min_section_heading(text):
    s = S("mH", fontName="Helvetica-Bold", fontSize=8.5,
           textColor=C_NAVY, leading=10, spaceBefore=10, spaceAfter=1)
    return KeepTogether([
        Paragraph(text.upper(), s),
        HRFlowable(width="100%", thickness=2, color=C_TEAL,
                   spaceAfter=3, spaceBefore=0),
    ])


def min_main_story(data, compact=False, max_bullets=None):
    s_job   = S("mJT", fontName="Helvetica-Bold", fontSize=9.8,
                 textColor=C_NAVY, leading=13)
    s_co    = S("mCo", fontName="Helvetica-Oblique", fontSize=8.5,
                 textColor=C_TEAL, leading=11, spaceAfter=1)
    s_dates = S("mDt", fontName="Helvetica-Bold", fontSize=8,
                 textColor=C_TEAL, leading=12, alignment=TA_RIGHT)
    s_bull  = S("mBu", fontName="Helvetica", fontSize=8.8,
                 textColor=C_GRAY, leading=12.5 if not compact else 11.5,
                 leftIndent=11, firstLineIndent=-8,
                 spaceAfter=2 if not compact else 1)
    s_sum   = S("mSu", fontName="Helvetica", fontSize=9,
                 textColor=C_GRAY, leading=13.5, spaceAfter=3)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.1 * inch)]
    job_idx = 0

    for sec in data["sections"]:
        if sec["name"] in SIDEBAR_SECTIONS: continue
        story.append(min_section_heading(sec["name"]))

        table_rows = []
        for ln in sec["items"]:
            ls = ln.strip()
            if not ls: continue
            if ls.startswith("|") and "---" not in ls:
                table_rows.append(ls); continue
            if table_rows:
                story.append(_build_table(table_rows, MIN_MAIN_W)); table_rows = []
            story.append(Paragraph(rich(ls), s_sum))
        if table_rows:
            story.append(_build_table(table_rows, MIN_MAIN_W))

        for job in sec["jobs"]:
            cap = (max_bullets or {}).get(job_idx)
            left  = Paragraph(f"<b>{job['title']}</b>", s_job)
            right = Paragraph(job["dates"], s_dates)
            t = Table([[left, right]],
                      colWidths=[MIN_MAIN_W * 0.68, MIN_MAIN_W * 0.32])
            t.setStyle(TableStyle([
                ("VALIGN",       (0,0),(-1,-1),"BOTTOM"),
                ("LEFTPADDING",  (0,0),(-1,-1),0),
                ("RIGHTPADDING", (0,0),(-1,-1),0),
                ("TOPPADDING",   (0,0),(-1,-1),0),
                ("BOTTOMPADDING",(0,0),(-1,-1),1),
            ]))
            block = [t]
            if job["company"]:
                block.append(Paragraph(job["company"], s_co))
            bullets = job["bullets"] if cap is None else job["bullets"][:cap]
            for b in bullets:
                block.append(Paragraph(f"–  {rich(b)}", s_bull))
            story.append(KeepTogether(block[:2]))
            for item in block[2:]: story.append(item)
            story.append(Spacer(1, 4 if not compact else 2))
            job_idx += 1

    return story


def make_min_frames():
    hf = Frame(0, H - MIN_HEADER_H, W, MIN_HEADER_H * 0.88,
               leftPadding=MIN_MX, rightPadding=MIN_MX * 0.3,
               topPadding=0, bottomPadding=0, id="hdr", showBoundary=0)
    sf = Frame(MIN_MX * 0.5, MIN_MY, MIN_SIDEBAR - MIN_MX * 0.5, MIN_CONTENT_H,
               leftPadding=10, rightPadding=6,
               topPadding=0, bottomPadding=0, id="side", showBoundary=0)
    mf = Frame(MIN_MAIN_X + MIN_MX * 0.3, MIN_MY, MIN_MAIN_W, MIN_CONTENT_H,
               leftPadding=3, rightPadding=3,
               topPadding=0, bottomPadding=0, id="main", showBoundary=0)
    return hf, sf, mf


# ══════════════════════════════════════════════════════════════════════════════
#  THEME 3 — BOLD ASYMMETRIC
# ══════════════════════════════════════════════════════════════════════════════
ASY_HEADER_H  = 1.15 * inch
ASY_SIDEBAR   = 2.2  * inch
ASY_MX        = 0.28 * inch
ASY_MY        = 0.28 * inch
ASY_MAIN_X    = ASY_SIDEBAR + 3
ASY_MAIN_W    = W - ASY_MAIN_X - ASY_MX * 0.6
ASY_CONTENT_H = H - ASY_HEADER_H - ASY_MY * 2


def draw_asymmetric(canvas, doc):
    canvas.saveState()
    # Header background (navy)
    canvas.setFillColor(C_NAVY)
    canvas.rect(0, H - ASY_HEADER_H, W, ASY_HEADER_H, fill=1, stroke=0)

    # Teal diagonal "slash" shape in header right side
    canvas.setFillColor(C_TEAL2)
    p = canvas.beginPath()
    p.moveTo(W * 0.48, H)
    p.lineTo(W * 0.60, H - ASY_HEADER_H)
    p.lineTo(W, H - ASY_HEADER_H)
    p.lineTo(W, H)
    p.close()
    canvas.drawPath(p, fill=1, stroke=0)

    # Gold accent slash layered on top (smaller, lighter)
    canvas.setFillColor(C_GOLD)
    p2 = canvas.beginPath()
    p2.moveTo(W * 0.44, H)
    p2.lineTo(W * 0.47, H - ASY_HEADER_H)
    p2.lineTo(W * 0.60, H - ASY_HEADER_H)
    p2.lineTo(W * 0.48, H)
    p2.close()
    canvas.drawPath(p2, fill=1, stroke=0)

    # Gold bottom border
    canvas.setFillColor(C_GOLD)
    canvas.rect(0, H - ASY_HEADER_H, W * 0.44, 3, fill=1, stroke=0)

    # Sidebar: gradient effect (3 stacked rects, darkest to darkest)
    for i, (y_start, col) in enumerate([
        (0,                 colors.HexColor("#0A1628")),
        (H * 0.3,           colors.HexColor("#0D1F38")),
        (H * 0.6,           colors.HexColor("#0F2340")),
    ]):
        canvas.setFillColor(col)
        canvas.rect(0, y_start, ASY_SIDEBAR,
                    (H * 0.3 if i < 2 else H * 0.4) - ASY_HEADER_H * (1 if i == 2 else 0),
                    fill=1, stroke=0)

    # Teal left accent strip
    canvas.setFillColor(C_TEAL)
    canvas.rect(0, 0, 4, H - ASY_HEADER_H, fill=1, stroke=0)

    # Gold rule between sidebar and main
    canvas.setFillColor(C_GOLD)
    canvas.rect(ASY_SIDEBAR, 0, 2.5, H - ASY_HEADER_H, fill=1, stroke=0)

    canvas.restoreState()


def asy_header_story(data):
    s_name = S("aN", fontName="Helvetica-Bold", fontSize=25,
                textColor=C_WHITE, leading=27, spaceAfter=2)
    s_sub  = S("aS", fontName="Helvetica", fontSize=9.5,
                textColor=C_GOLD2, leading=13, spaceAfter=0)
    s_cont = S("aC", fontName="Helvetica", fontSize=7.8,
                textColor=colors.HexColor("#9FBCD4"), leading=11)

    contact_parts = [p.strip() for p in data["contact"].split("|") if p.strip()]

    story = [Spacer(1, 0.12 * inch)]
    story.append(Paragraph(data["name"], s_name))
    if data["subtitle"]:
        story.append(Paragraph(data["subtitle"], s_sub))
    story.append(Spacer(1, 3))
    story.append(Paragraph("  ·  ".join(contact_parts), s_cont))
    return story


def asy_sidebar_story(data, compact=False):
    s_sec  = S("aSec", fontName="Helvetica-Bold", fontSize=7,
                textColor=C_GOLD, leading=9, spaceAfter=4, spaceBefore=10,
                letterSpacing=0.9)
    s_body = S("aBd",  fontName="Helvetica", fontSize=7.6,
                textColor=colors.HexColor("#9FBCD4"), leading=11, spaceAfter=1)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.14 * inch)]

    for sec in data["sections"]:
        sname = sec["name"]
        if sname not in SIDEBAR_SECTIONS: continue

        story.append(Paragraph(sname.upper(), s_sec))
        story.append(HRFlowable(width="100%", thickness=0.8, color=C_GOLD,
                                spaceAfter=3, spaceBefore=0))

        if sname in ("Core Skills", "Technical Skills", "Skills"):
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    sk = strip_md(ls[2:]).strip()
                    if len(sk) > 42: sk = sk[:40] + "…"
                    story.append(SkillBar(sk, ASY_SIDEBAR - 20,
                                         fill=0.80, bar_color=C_GOLD))
                elif ls.startswith("|") and "---" not in ls:
                    cells = [c.strip() for c in ls.split("|")[1:-1]]
                    if cells and strip_md(cells[0]).lower() not in ("category", ""):
                        cat  = strip_md(cells[0])
                        vals = strip_md(cells[1]) if len(cells) > 1 else ""
                        story.append(Paragraph(f"<b>{cat}</b>", s_body))
                        for v in vals.split(","):
                            v = v.strip()
                            if v:
                                story.append(SkillBar(v, ASY_SIDEBAR - 20,
                                                      fill=0.80, bar_color=C_GOLD))

        elif sname == "Education":
            for ln in sec["items"]:
                ls = ln.strip()
                if not ls or ls.startswith("|"): continue
                story.append(Paragraph(strip_md(ls), s_body))

        elif sname == "Certifications":
            for ln in sec["items"]:
                ls = ln.strip()
                if ls.startswith("- "):
                    story.append(Paragraph(f"◆  {strip_md(ls[2:])}", s_body))

    return story


def asy_section_heading(text):
    return KeepTogether([
        GoldBarHeading(text, ASY_MAIN_W, bar_color=C_TEAL, text_color=C_NAVY),
        Spacer(1, 2),
    ])


def asy_main_story(data, compact=False, max_bullets=None):
    s_job   = S("aJT", fontName="Helvetica-Bold", fontSize=10,
                 textColor=C_NAVY, leading=13)
    s_co    = S("aCo", fontName="Helvetica-Oblique", fontSize=8.5,
                 textColor=C_TEAL, leading=11, spaceAfter=1)
    s_dates = S("aDt", fontName="Helvetica-Bold", fontSize=8,
                 textColor=C_GOLD, leading=12, alignment=TA_RIGHT)
    s_bull  = S("aBu", fontName="Helvetica", fontSize=8.8,
                 textColor=C_GRAY, leading=12.5 if not compact else 11.5,
                 leftIndent=11, firstLineIndent=-8,
                 spaceAfter=2 if not compact else 1)
    s_sum   = S("aSu", fontName="Helvetica", fontSize=9,
                 textColor=C_GRAY, leading=13.5, spaceAfter=3)

    SIDEBAR_SECTIONS = {"Core Skills", "Technical Skills", "Skills",
                        "Education", "Certifications"}
    story = [Spacer(1, 0.1 * inch)]
    job_idx = 0

    for sec in data["sections"]:
        if sec["name"] in SIDEBAR_SECTIONS: continue
        story.append(asy_section_heading(sec["name"]))

        table_rows = []
        for ln in sec["items"]:
            ls = ln.strip()
            if not ls: continue
            if ls.startswith("|") and "---" not in ls:
                table_rows.append(ls); continue
            if table_rows:
                story.append(_build_table(table_rows, ASY_MAIN_W)); table_rows = []
            story.append(Paragraph(rich(ls), s_sum))
        if table_rows:
            story.append(_build_table(table_rows, ASY_MAIN_W))

        for job in sec["jobs"]:
            cap = (max_bullets or {}).get(job_idx)
            left  = Paragraph(f"<b>{job['title']}</b>", s_job)
            right = Paragraph(job["dates"], s_dates)
            t = Table([[left, right]],
                      colWidths=[ASY_MAIN_W * 0.70, ASY_MAIN_W * 0.30])
            t.setStyle(TableStyle([
                ("VALIGN",       (0,0),(-1,-1),"BOTTOM"),
                ("LEFTPADDING",  (0,0),(-1,-1),0),
                ("RIGHTPADDING", (0,0),(-1,-1),0),
                ("TOPPADDING",   (0,0),(-1,-1),0),
                ("BOTTOMPADDING",(0,0),(-1,-1),1),
            ]))
            block = [t]
            if job["company"]:
                block.append(Paragraph(job["company"], s_co))
            bullets = job["bullets"] if cap is None else job["bullets"][:cap]
            for b in bullets:
                block.append(Paragraph(
                    f"<font color='#C9A84C'>◆</font>  {rich(b)}", s_bull))
            story.append(KeepTogether(block[:2]))
            for item in block[2:]: story.append(item)
            story.append(Spacer(1, 3 if not compact else 1))
            job_idx += 1

    return story


def make_asy_frames():
    hf = Frame(0, H - ASY_HEADER_H, W * 0.90, ASY_HEADER_H * 0.90,
               leftPadding=ASY_MX, rightPadding=ASY_MX * 0.3,
               topPadding=0, bottomPadding=0, id="hdr", showBoundary=0)
    sf = Frame(ASY_MX * 0.4, ASY_MY, ASY_SIDEBAR - ASY_MX * 0.4, ASY_CONTENT_H,
               leftPadding=12, rightPadding=6,
               topPadding=0, bottomPadding=0, id="side", showBoundary=0)
    mf = Frame(ASY_MAIN_X + ASY_MX * 0.2, ASY_MY, ASY_MAIN_W, ASY_CONTENT_H,
               leftPadding=3, rightPadding=3,
               topPadding=0, bottomPadding=0, id="main", showBoundary=0)
    return hf, sf, mf


# ── Shared table builder ──────────────────────────────────────────────────────
def _build_table(rows_raw, main_w):
    rows = []
    for ln in rows_raw:
        cells = [c.strip() for c in ln.split("|")[1:-1]]
        if not cells: continue
        styled = []
        for c_i, cell in enumerate(cells):
            cell = strip_md(cell)
            st = ParagraphStyle("TC", fontName=(
                "Helvetica-Bold" if c_i == 0 else "Helvetica"),
                fontSize=8.5, leading=12,
                textColor=(C_TEAL if c_i == 0 else C_GRAY))
            styled.append(Paragraph(cell, st))
        rows.append(styled)
    if not rows:
        return Spacer(1, 0)
    t = Table(rows, colWidths=[main_w * 0.34, main_w * 0.66])
    t.setStyle(TableStyle([
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [C_LTEAL, C_WHITE]),
        ("BOX",           (0,0),(-1,-1), 0.5, C_RULE),
        ("INNERGRID",     (0,0),(-1,-1), 0.25, C_RULE),
        ("VALIGN",        (0,0),(-1,-1), "TOP"),
        ("LEFTPADDING",   (0,0),(-1,-1), 5),
        ("RIGHTPADDING",  (0,0),(-1,-1), 5),
        ("TOPPADDING",    (0,0),(-1,-1), 3),
        ("BOTTOMPADDING", (0,0),(-1,-1), 3),
    ]))
    return t


# ── Auto-fit ──────────────────────────────────────────────────────────────────
def test_pages(hdr, sidebar, main, frames, draw_fn):
    buf = BytesIO()
    hf, sf, mf = frames()
    doc = BaseDocTemplate(buf, pagesize=letter,
                          pageTemplates=[PageTemplate(
                              id="T", frames=[hf, sf, mf], onPage=draw_fn)],
                          leftMargin=0, rightMargin=0, topMargin=0, bottomMargin=0)
    doc.build(hdr + [FrameBreak()] + sidebar + [FrameBreak()] + main)
    return doc.page


def auto_fit(data, theme):
    fns = THEMES[theme]
    hdr_fn, side_fn, main_fn, frames_fn, draw_fn = fns

    hdr = hdr_fn(data)
    for compact in (False, True):
        side = side_fn(data, compact=compact)
        main = main_fn(data, compact=compact)
        if test_pages(hdr, side, main, frames_fn, draw_fn) == 1:
            return compact, None

    # Trim bullets from oldest jobs until it fits
    n_jobs = sum(len(s.get("jobs", [])) for s in data["sections"])
    max_b = {}
    idx = 0
    for sec in data["sections"]:
        for job in sec.get("jobs", []):
            max_b[idx] = len(job["bullets"])
            idx += 1

    for _ in range(60):
        for ji in range(n_jobs - 1, -1, -1):
            if max_b.get(ji, 0) > 0:
                max_b[ji] -= 1
                break
        side = side_fn(data, compact=True)
        main = main_fn(data, compact=True, max_bullets=max_b)
        if test_pages(hdr, side, main, frames_fn, draw_fn) == 1:
            return True, max_b

    return True, max_b


# ── Theme registry ────────────────────────────────────────────────────────────
THEMES = {
    "executive":  (exec_header_story, exec_sidebar_story, exec_main_story,
                   make_exec_frames, draw_executive),
    "minimal":    (min_header_story,  min_sidebar_story,  min_main_story,
                   make_min_frames,   draw_minimal),
    "asymmetric": (asy_header_story,  asy_sidebar_story,  asy_main_story,
                   make_asy_frames,   draw_asymmetric),
}


# ── PDF builder ───────────────────────────────────────────────────────────────
def build_pdf_resume(md: str, out_path: str, style: str = "executive"):
    data = parse_resume(md)
    hdr_fn, side_fn, main_fn, frames_fn, draw_fn = THEMES[style]

    compact, max_b = auto_fit(data, style)
    hdr  = hdr_fn(data)
    side = side_fn(data, compact=compact)
    main = main_fn(data, compact=compact, max_bullets=max_b)

    hf, sf, mf = frames_fn()
    doc = BaseDocTemplate(
        out_path, pagesize=letter,
        pageTemplates=[PageTemplate(id="T", frames=[hf, sf, mf], onPage=draw_fn)],
        leftMargin=0, rightMargin=0, topMargin=0, bottomMargin=0
    )
    doc.build(hdr + [FrameBreak()] + side + [FrameBreak()] + main)
    pages = doc.page
    tag = "(1p)" if pages == 1 else f"({pages}p ⚠️)"
    print(f"✓ PDF {tag} [{style}]: {out_path}")


# ── Word Cover Letter ─────────────────────────────────────────────────────────
def build_word_cover_letter(md: str, out_path: str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_para(text="", bold=False, size=11, color=None,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if not text: return p
        for part in re.split(r'(\*\*.*?\*\*)', text):
            is_b = part.startswith("**") and part.endswith("**")
            run = p.add_run(part[2:-2] if is_b else part)
            run.bold = bold or is_b
            run.font.size = Pt(size)
            run.font.name = "Calibri"
            if color: run.font.color.rgb = color
        return p

    def add_rule(color_hex="1B7A8C", thickness=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    str(thickness))
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    lines = md.strip().split("\n")
    i = 0
    while i < len(lines) and not lines[i].strip().startswith("# "): i += 1
    i += 1

    name_line = lines[i].strip() if i < len(lines) else ""
    add_para(name_line, bold=True, size=18,
             color=RGBColor(0x0D, 0x1B, 0x2A), space_after=2)
    i += 1

    while i < len(lines) and lines[i].strip():
        add_para(lines[i].strip(), size=10,
                 color=RGBColor(0x71, 0x80, 0x96), space_after=0)
        i += 1
    add_rule("1B7A8C", 16)

    while i < len(lines) and not lines[i].strip(): i += 1
    while i < len(lines) and lines[i].strip() \
          and not lines[i].strip().startswith("Dear"):
        add_para(lines[i].strip(), size=10.5,
                 color=RGBColor(0x0D, 0x1B, 0x2A), space_after=0)
        i += 1
    add_para()

    while i < len(lines):
        ln = lines[i].strip()
        if not ln:
            add_para(space_after=4); i += 1; continue
        if ln.startswith("Dear"):
            add_para(ln, bold=True, size=11,
                     color=RGBColor(0x0D, 0x1B, 0x2A), space_after=10)
            i += 1; continue
        if ln.startswith("Sincerely"):
            add_para()
            add_para(ln, size=11, color=RGBColor(0x0D, 0x1B, 0x2A), space_after=0)
            i += 1; continue
        add_para(ln, size=11, color=RGBColor(0x0D, 0x1B, 0x2A), space_after=8)
        i += 1

    doc.save(out_path)
    print(f"✓ DOCX: {out_path}")


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    args = sys.argv[1:]
    if not args:
        print("Usage: python3 generate_application.py <folder> [--style executive|minimal|asymmetric]")
        sys.exit(1)

    folder = args[0]
    style  = "executive"
    if "--style" in args:
        idx = args.index("--style")
        if idx + 1 < len(args):
            style = args[idx + 1]

    if style not in THEMES:
        print(f"Unknown style '{style}'. Choose: {', '.join(THEMES)}")
        sys.exit(1)

    with open(os.path.join(folder, "resume.md"))       as f: resume_md = f.read()
    with open(os.path.join(folder, "cover_letter.md")) as f: cover_md  = f.read()

    build_pdf_resume(resume_md,       os.path.join(folder, "resume.pdf"), style=style)
    build_word_cover_letter(cover_md, os.path.join(folder, "cover_letter.docx"))
